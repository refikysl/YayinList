import streamlit as st
import db_manager
import apa_formatter
import bibtex_helper
from datetime import date
import os

# Turkish-aware Title Case function
def turkish_title_case(text):
    """
    Convert text to Title Case with Turkish character support.
    Handles: İ/i, I/ı properly
    Preserves: Abbreviations like "ve", "için", "ile"
    """
    if not text or not isinstance(text, str):
        return text
    
    # Words that should remain lowercase (Turkish articles, conjunctions, prepositions)
    lowercase_words = {'ve', 'veya', 'ile', 'için', 'de', 'da', 'den', 'dan', 'bir', 'bu', 'şu', 'o'}
    
    words = text.split()
    result = []
    
    for i, word in enumerate(words):
        # First word or not a lowercase exception
        if i == 0 or word.lower() not in lowercase_words:
            # Turkish-aware capitalize
            if word:
                # Handle Turkish i/İ properly
                first_char = word[0]
                if first_char == 'i':
                    capitalized = 'İ' + word[1:].lower()
                elif first_char == 'ı':
                    capitalized = 'I' + word[1:].lower()
                else:
                    capitalized = first_char.upper() + word[1:].lower()
                result.append(capitalized)
            else:
                result.append(word)
        else:
            result.append(word.lower())
    
    return ' '.join(result)


# Page configuration
st.set_page_config(page_title="Akademik Yayın Yönetim Sistemi", page_icon="📚", layout="wide")

# Check connectivity
if not db_manager.get_api_url():
    st.warning("⚠️ Sistem henüz kurulmamış.")
    st.info("Lütfen kurulum kılavuzundaki adımları takip ederek aldığınız **Web App URL**'sini 'api_url.txt' dosyasına kaydedin.")
    
    url_input = st.text_input("Veya URL'yi buraya yapıştırıp Enter'a basın:", placeholder="https://script.google.com/macros/s/...")
    if url_input:
        with open("api_url.txt", "w") as f:
            f.write(url_input)
        st.success("URL kaydedildi! Sayfa yenileniyor...")
        st.rerun()
    st.stop()

# Initialize DB
db_manager.init_db()

# Sidebar
st.sidebar.title("Navigasyon")
page = st.sidebar.radio("Sayfa Seçiniz:", ["Veri Girişi (Hocalar/Asistanlar)", "Raporlama (Admin)"])

# --- Session State for Dynamic Rows ---
if 'num_authors' not in st.session_state:
    st.session_state.num_authors = 1
if 'num_editors' not in st.session_state:
    st.session_state.num_editors = 1

# --- Pre-fill State Helper ---
def prefill_state(key, value):
    if value:
        st.session_state[key] = value

if page == "Veri Girişi (Hocalar/Asistanlar)":
    st.title("📚 Yayın Veri Girişi")
    
    # Session State
    if 'reset_counter' not in st.session_state:
        st.session_state.reset_counter = 0
    rk = str(st.session_state.reset_counter)
    
    # Success message (compact)
    if 'success_msg' in st.session_state:
        st.success(st.session_state.success_msg)
    
    # My Publications Search Section
    with st.expander("🔍 Kayıtlarımı Görüntüle", expanded=False):
        st.markdown("**Daha önce girdiğiniz yayınları görmek için soyadınızı yazın:**")
        
        col_search1, col_search2 = st.columns([3, 1])
        with col_search1:
            search_surname = st.text_input(
                "Soyad",
                placeholder="Örn: Yılmaz",
                key="search_surname_input",
                label_visibility="collapsed"
            )
        with col_search2:
            search_button = st.button("🔎 Ara", use_container_width=True)
        
        if search_button and search_surname:
            # Search publications by surname
            all_pubs = db_manager.get_all_publications()
            
            if all_pubs:
                # Filter by surname (case-insensitive)
                search_lower = search_surname.lower().strip()
                matching_pubs = []
                
                for pub in all_pubs:
                    authors = pub.get('authors', [])
                    if isinstance(authors, list):
                        for author in authors:
                            surname = author.get('surname', '').lower().strip()
                            if search_lower in surname or surname in search_lower:
                                matching_pubs.append(pub)
                                break
                
                if matching_pubs:
                    st.success(f"✅ {len(matching_pubs)} yayın bulundu:")
                    
                    # Group by type
                    grouped = {}
                    for pub in matching_pubs:
                        ptype = pub.get('publication_type', 'Diğer')
                        if ptype not in grouped:
                            grouped[ptype] = []
                        grouped[ptype].append(pub)
                    
                    # Display grouped publications
                    for ptype in ["Makale", "Kitap", "Kitap Bölümü", "Bildiri", "Proje", "Diğer"]:
                        if ptype in grouped:
                            st.markdown(f"**{ptype} ({len(grouped[ptype])})**")
                            for idx, pub in enumerate(grouped[ptype], 1):
                                citation = apa_formatter.format_apa_6(pub)
                                st.markdown(f"{idx}. {citation}")
                            st.markdown("")
                else:
                    st.warning(f"'{search_surname}' soyadıyla kayıt bulunamadı.")
            else:
                st.info("Henüz hiç yayın kaydedilmemiş.")
        elif search_button and not search_surname:
            st.warning("Lütfen soyadınızı girin.")
    
    st.markdown("---")

    
    # Top row: Department, Type, New Button
    col_t1, col_t2, col_t3 = st.columns([2, 2, 1])
    
    with col_t1:
        department = st.selectbox(
            "Bölüm",
            [
                "Siyaset Bilimi ve Kamu Yönetimi",
                "İktisat",
                "İşletme",
                "Maliye",
                "Ekonometri",
                "Uluslararası İlişkiler"
            ]
        )
    
    with col_t2:
        # Handle BibTeX publication type
        pub_types = ["Makale", "Kitap", "Kitap Bölümü", "Bildiri", "Proje"]
        
        # If BibTeX type exists, set it in session state for the widget
        if 'bibtex_pub_type' in st.session_state:
            bibtex_type = st.session_state['bibtex_pub_type']
            if bibtex_type in pub_types:
                # Set the widget's session state value directly
                st.session_state['pub_type_selectbox'] = bibtex_type
            # Clear the flag
            del st.session_state['bibtex_pub_type']
        
        # Initialize if not present
        if 'pub_type_selectbox' not in st.session_state:
            st.session_state['pub_type_selectbox'] = "Makale"
        
        pub_type = st.selectbox(
            "Yayın Türü",
            pub_types,
            key='pub_type_selectbox'
        )
    
    with col_t3:
        if st.button("🆕 Yeni", type="secondary", use_container_width=True, help="Formu temizle"):
            # Increment reset counter
            st.session_state.reset_counter += 1
            
            # Clear ALL form-related session state keys
            keys_to_delete = []
            for key in list(st.session_state.keys()):
                # Comprehensive pattern matching
                if any(pattern in key for pattern in [
                    'input_', 'surname_', 'name_', 'ed_surname_', 'ed_name_',
                    'last_bib_file', 'show_date_msg', 'success_msg', 
                    'bibtex_pub_type', 'pub_type_selectbox', 'bibtex_uploader_',
                    'del_auth_', 'del_ed_', 'add_auth_', 'add_ed_'
                ]):
                    keys_to_delete.append(key)
            
            for key in keys_to_delete:
                del st.session_state[key]
            
            # Reset counters
            st.session_state.num_authors = 1
            st.session_state.num_editors = 1
            
            st.rerun()
    
    # BibTeX Upload (Compact Expander)
    with st.expander("📂 BibTeX Yükle", expanded=False):
        # Use reset_counter in key to force recreation
        uploaded_file = st.file_uploader("", type=['bib'], label_visibility="collapsed", key=f'bibtex_uploader_{rk}')
        
        if uploaded_file is None:
            if 'last_bib_file' in st.session_state:
                del st.session_state.last_bib_file
        else:
            file_signature = f"{uploaded_file.name}_{uploaded_file.size}"
            if 'last_bib_file' not in st.session_state or st.session_state.last_bib_file != file_signature:
                string_data = uploaded_file.getvalue().decode("utf-8")
                parsed = bibtex_helper.parse_bibtex(string_data)
                
                if parsed:
                    st.session_state.last_bib_file = file_signature
                    st.success("✅ Yüklendi!")
                    
                    # Store publication type separately (will be used in selectbox index)
                    if 'publication_type' in parsed:
                        st.session_state['bibtex_pub_type'] = parsed['publication_type']
                    
                    prefill_state(f'title_input_{rk}', parsed.get('title'))
                    
                    if 'publication_date' in parsed:
                        st.session_state[f'pub_date_input_{rk}'] = parsed['publication_date']
                        st.session_state['show_date_msg'] = True
                    
                    if 'authors' in parsed and parsed['authors']:
                        auths = parsed['authors']
                        st.session_state.num_authors = len(auths)
                        for idx, auth in enumerate(auths, 1):
                            prefill_state(f"surname_{idx}_{rk}", auth.get('surname'))
                            prefill_state(f"name_{idx}_{rk}", auth.get('name'))
                    
                    if 'editors' in parsed and parsed['editors']:
                        eds = parsed['editors']
                        st.session_state.num_editors = len(eds)
                        for idx, ed in enumerate(eds, 1):
                            prefill_state(f"ed_surname_{idx}_{rk}", ed.get('surname'))
                            prefill_state(f"ed_name_{idx}_{rk}", ed.get('name'))
                    
                    prefill_state(f'journal_name_input_{rk}', parsed.get('journal_name'))
                    prefill_state(f'volume_input_{rk}', parsed.get('volume'))
                    prefill_state(f'issue_input_{rk}', parsed.get('issue'))
                    prefill_state(f'pages_input_{rk}', parsed.get('pages'))
                    prefill_state(f'publisher_input_{rk}', parsed.get('publisher'))
                    prefill_state(f'location_input_{rk}', parsed.get('location'))
                    prefill_state(f'book_title_input_{rk}', parsed.get('book_title'))
                    
                    st.rerun()
                else:
                    st.warning("Dosya okunamadı")
    
    # Main form in 2 columns
    col_left, col_right = st.columns([1, 1])
    
    with col_left:
        # Title and Date
        title = st.text_input("📝 Başlık", key=f'title_input_{rk}')
        
        # Initialize date in session state if not present
        if f'pub_date_input_{rk}' not in st.session_state:
            st.session_state[f'pub_date_input_{rk}'] = date.today()
        
        publication_date = st.date_input("📅 Tarih", key=f'pub_date_input_{rk}')
        
        if st.session_state.get('show_date_msg'):
            st.info("📅 BibTeX kaydında yayının tarih bilgisi yalnızca yıl olarak yer aldığından o yılın 1 Ocak tarihi olarak belirlenmiştir. Yukarıdaki alandan tarih bilgisini gün ve ayı içerecek şekilde değiştirebilirsiniz.")
        
        # Authors (Compact)
        st.markdown("**👥 Yazarlar**")
        authors_data = []
        for i in range(1, st.session_state.num_authors + 1):
            col_a1, col_a2, col_a3 = st.columns([2, 2, 0.5])
            with col_a1:
                surname = st.text_input("Soyad", key=f"surname_{i}_{rk}", label_visibility="collapsed", placeholder=f"{i}. Soyadı")
            with col_a2:
                name = st.text_input("Ad", key=f"name_{i}_{rk}", label_visibility="collapsed", placeholder=f"{i}. Adı")
            with col_a3:
                if i > 1 and st.button("✖", key=f"del_auth_{i}_{rk}", help="Çıkar"):
                    st.session_state.num_authors -= 1
                    st.rerun()
            
            if surname and name:
                authors_data.append({'surname': surname, 'name': name})
        
        if st.session_state.num_authors < 5:
            if st.button("➕ Yazar Ekle", key=f"add_auth_{rk}"):
                st.session_state.num_authors += 1
                st.rerun()
    
    with col_right:
        # Type-specific fields
        data = {}
        missing_fields = []
        editors_data = []
        
        st.markdown(f"**📋 {pub_type} Bilgileri**")
        
        if pub_type == "Makale":
            journal_name = st.text_input("Dergi", key=f'journal_name_input_{rk}')
            col_v1, col_v2, col_v3 = st.columns(3)
            with col_v1:
                volume = st.text_input("Cilt", key=f'volume_input_{rk}')
            with col_v2:
                issue = st.text_input("Sayı", key=f'issue_input_{rk}')
            with col_v3:
                pages = st.text_input("Sayfa", key=f'pages_input_{rk}')
            
            if not journal_name: missing_fields.append("Dergi")
            data.update({'journal_name': journal_name, 'volume': volume, 'issue': issue, 'pages': pages})
            
        elif pub_type == "Kitap":
            publisher = st.text_input("Yayınevi", key=f'publisher_input_{rk}')
            location = st.text_input("Basım Yeri", key=f'location_input_{rk}')
            
            if not publisher: missing_fields.append("Yayınevi")
            if not location: missing_fields.append("Basım Yeri")
            data.update({'publisher': publisher, 'location': location})
            
        elif pub_type == "Kitap Bölümü":
            book_title = st.text_input("Kitap Adı", key=f'book_title_input_{rk}')
            col_p1, col_p2 = st.columns(2)
            with col_p1:
                publisher = st.text_input("Yayınevi", key=f'publisher_input_{rk}')
            with col_p2:
                location = st.text_input("Basım Yeri", key=f'location_input_{rk}')
            pages = st.text_input("Sayfa", key=f'pages_input_{rk}')
            
            # Editors in expander
            with st.expander("✏️ Editörler"):
                for j in range(1, st.session_state.num_editors + 1):
                    col_e1, col_e2, col_e3 = st.columns([2, 2, 0.5])
                    with col_e1:
                        e_surname = st.text_input("Soyad", key=f"ed_surname_{j}_{rk}", label_visibility="collapsed", placeholder=f"{j}. Ed. Soyadı")
                    with col_e2:
                        e_name = st.text_input("Ad", key=f"ed_name_{j}_{rk}", label_visibility="collapsed", placeholder=f"{j}. Ed. Adı")
                    with col_e3:
                        if j > 1 and st.button("✖", key=f"del_ed_{j}_{rk}", help="Çıkar"):
                            st.session_state.num_editors -= 1
                            st.rerun()
                    
                    if e_surname and e_name:
                        editors_data.append({'surname': e_surname, 'name': e_name})
                
                if st.session_state.num_editors < 5:
                    if st.button("➕ Editör Ekle", key=f"add_ed_{rk}"):
                        st.session_state.num_editors += 1
                        st.rerun()
            
            if not book_title: missing_fields.append("Kitap Adı")
            if not publisher: missing_fields.append("Yayınevi")
            
            data.update({
                'book_title': book_title,
                'publisher': publisher,
                'location': location,
                'pages': pages,
                'editors': editors_data
            })
            
        elif pub_type == "Bildiri":
            conf_name = st.text_input("Konferans", key=f'book_title_input_{rk}')
            col_c1, col_c2 = st.columns(2)
            with col_c1:
                location = st.text_input("Yer", key=f'location_input_{rk}')
            with col_c2:
                publisher = st.text_input("Organizasyon", key=f'publisher_input_{rk}')
            
            if not conf_name: missing_fields.append("Konferans")
            data.update({'book_title': conf_name, 'location': location, 'publisher': publisher})
            
        elif pub_type == "Proje":
            funding_agency = st.text_input("Destekleyen Kurum", key=f'funding_agency_input_{rk}')
            project_status = st.text_input("Proje No", key=f'project_status_input_{rk}')
            if not funding_agency: missing_fields.append("Kurum")
            data.update({'funding_agency': funding_agency, 'project_status': project_status})
    
    # Save button at bottom
    st.markdown("---")
    
    # Show success message here too
    if 'success_msg' in st.session_state:
        st.success(st.session_state.success_msg)
    
    submitted = st.button("💾 Yayını Kaydet", type="primary")
    
    if submitted:
        if not authors_data:
            st.error("En az 1 yazar girmelisiniz.")
        elif pub_type == "Kitap Bölümü" and not editors_data:
             st.error("En az 1 editör girmelisiniz.")
        elif not title:
            st.error("Başlık girmelisiniz.")
        elif missing_fields:
            st.error(f"Eksik alanlar: {', '.join(missing_fields)}")
        else:
            # Normalize all text fields to Title Case
            # Authors
            normalized_authors = []
            for author in authors_data:
                normalized_authors.append({
                    'name': turkish_title_case(author.get('name', '')),
                    'surname': turkish_title_case(author.get('surname', ''))
                })
            
            # Editors (if any)
            normalized_editors = []
            if 'editors' in data and data['editors']:
                for editor in data['editors']:
                    normalized_editors.append({
                        'name': turkish_title_case(editor.get('name', '')),
                        'surname': turkish_title_case(editor.get('surname', ''))
                    })
            
            # Normalize other text fields
            normalized_data = {}
            for key, value in data.items():
                if key == 'editors':
                    normalized_data[key] = normalized_editors
                elif isinstance(value, str) and key not in ['volume', 'issue', 'pages']:
                    # Apply title case to text fields (not numbers)
                    normalized_data[key] = turkish_title_case(value)
                else:
                    normalized_data[key] = value
            
            full_data = {
                'department': department,
                'publication_type': pub_type,
                'authors': normalized_authors,
                'publication_date': publication_date.strftime("%Y-%m-%d"),
                'title': turkish_title_case(title),
                **normalized_data
            }
            
            success = db_manager.add_publication(full_data)
            
            if success:
                st.toast("✅ Yayın başarıyla kaydedildi!", icon="✅")
                apa_citation = apa_formatter.format_apa_6(full_data)
                st.session_state.success_msg = f"**{pub_type} başarıyla veritabanına kaydedildi!**\n\n**APA Formatı:** {apa_citation}"
                st.success(st.session_state.success_msg)
                st.info("💡 Yeni yayın eklemek için yukarıdaki '🆕 Yeni' butonuna basın.")

elif page == "Raporlama (Admin)":
    st.title("📊 Yönetici Rapor Ekranı")
    
    if 'admin_unlocked' not in st.session_state:
        st.session_state.admin_unlocked = False
        
    if not st.session_state.admin_unlocked:
        password = st.text_input("Yönetici Şifresi", type="password")
        if st.button("Giriş Yap"):
            if password == "1379":
                st.session_state.admin_unlocked = True
                st.rerun()
            else:
                st.error("Hatalı Şifre!")
    else:
        if st.button("Çıkış Yap"):
            st.session_state.admin_unlocked = False
            st.rerun()
            
        st.markdown("---")
        st.markdown("### Rapor Filtreleme")
        
        # Department Selection
        st.markdown("#### 🏛️ Bölüm Seçimi")
        selected_department = st.selectbox(
            "Bölüm Seçiniz:",
            [
                "Tümü",
                "Siyaset Bilimi ve Kamu Yönetimi",
                "İktisat",
                "İşletme",
                "Maliye",
                "Ekonometri",
                "Uluslararası İlişkiler"
            ]
        )
        
        st.markdown("---")
        
        # Report Type Selection
        report_type = st.selectbox(
            "Raporlama Türü",
            ["Bölüm ve Tür Bazında Detaylı Rapor", "Tüm Yayınlar", "Yayın Türü Bazında", "Kişi Bazında"]
        )
        
        col1, col2 = st.columns(2)
        with col1:
            start_date = st.date_input("Başlangıç Tarihi", value=date(2025, 1, 1))
        with col2:
            end_date = st.date_input("Bitiş Tarihi", value=date(2026, 1, 1))
        
        # Additional filters based on report type
        selected_pub_type = None
        selected_person = None
        
        if report_type == "Yayın Türü Bazında":
            selected_pub_type = st.selectbox(
                "Yayın Türü Seçiniz",
                ["Makale", "Kitap", "Kitap Bölümü", "Bildiri", "Proje"]
            )
        elif report_type == "Kişi Bazında":
            st.info("Kişi adını tam olarak yazınız (Örn: Yılmaz)")
            selected_person = st.text_input("Yazar Soyadı")
            
        if st.button("Raporu Getir"):
            s_date_str = start_date.strftime("%Y-%m-%d")
            e_date_str = end_date.strftime("%Y-%m-%d")
            
            with st.spinner("Google Sheets'ten veriler çekiliyor..."):
                publications = db_manager.get_publications(s_date_str, e_date_str)
            
            if not publications:
                st.warning("Bu tarih aralığında yayın bulunamadı.")
            else:
                # Apply department filter first
                if selected_department != "Tümü":
                    publications = [p for p in publications if p.get('department') == selected_department]
                
                if not publications:
                    st.warning(f"{selected_department} bölümünde bu tarih aralığında yayın bulunamadı.")
                else:
                    # Apply additional filters
                    filtered_pubs = publications
                    
                    if report_type == "Yayın Türü Bazında" and selected_pub_type:
                        filtered_pubs = [p for p in publications if p.get('publication_type') == selected_pub_type]
                        
                    elif report_type == "Kişi Bazında" and selected_person:
                        # Filter by author surname
                        person_pubs = []
                        for p in publications:
                            authors = p.get('authors', [])
                            if isinstance(authors, list):
                                for auth in authors:
                                    if isinstance(auth, dict):
                                        surname = auth.get('surname', '').strip().lower()
                                        if selected_person.strip().lower() in surname:
                                            person_pubs.append(p)
                                            break
                        filtered_pubs = person_pubs
                    
                    if not filtered_pubs:
                        st.warning("Seçilen kriterlere uygun yayın bulunamadı.")
                    else:
                        # NEW: Department and Type Detailed Report
                        if report_type == "Bölüm ve Tür Bazında Detaylı Rapor":
                            st.subheader(f"📊 Detaylı Rapor - Toplam {len(filtered_pubs)} Yayın")
                            
                            # Group by department first, then by type
                            dept_groups = {}
                            for pub in filtered_pubs:
                                dept = pub.get('department', 'Belirtilmemiş')
                                if dept not in dept_groups:
                                    dept_groups[dept] = {}
                                
                                ptype = pub.get('publication_type', 'Diğer')
                                if ptype not in dept_groups[dept]:
                                    dept_groups[dept][ptype] = []
                                dept_groups[dept][ptype].append(pub)
                            
                            report_text = ""
                            
                            # Sort departments
                            dept_order = [
                                "Siyaset Bilimi ve Kamu Yönetimi",
                                "İktisat",
                                "İşletme",
                                "Maliye",
                                "Ekonometri",
                                "Uluslararası İlişkiler",
                                "Belirtilmemiş"
                            ]
                            
                            for dept in dept_order:
                                if dept in dept_groups:
                                    dept_total = sum(len(pubs) for pubs in dept_groups[dept].values())
                                    st.markdown(f"### 🏛️ {dept} ({dept_total} yayın)")
                                    report_text += f"\n## {dept} ({dept_total} yayın)\n\n"
                                    
                                    # Sort by publication type
                                    for ptype in ["Makale", "Kitap", "Kitap Bölümü", "Bildiri", "Proje", "Diğer"]:
                                        if ptype in dept_groups[dept]:
                                            pubs = dept_groups[dept][ptype]
                                            st.markdown(f"#### 📄 {ptype} ({len(pubs)})")
                                            report_text += f"\n### {ptype} ({len(pubs)})\n\n"
                                            
                                            for idx, pub in enumerate(pubs, 1):
                                                citation = apa_formatter.format_apa_6(pub)
                                                st.markdown(f"**{idx}.** {citation}")
                                                # Keep italics for export
                                                report_text += f"{idx}. {citation}\n\n"
                                            
                                            st.markdown("")  # Add spacing
                                    
                                    st.markdown("---")
                        
                        # Group by publication type if "Tüm Yayınlar"
                        elif report_type == "Tüm Yayınlar":
                            st.subheader(f"Bulunan Yayınlar ({len(filtered_pubs)})")
                            
                            # Group by type
                            grouped = {}
                            for pub in filtered_pubs:
                                ptype = pub.get('publication_type', 'Diğer')
                                if ptype not in grouped:
                                    grouped[ptype] = []
                                grouped[ptype].append(pub)
                            
                            report_text = ""
                            for ptype in ["Makale", "Kitap", "Kitap Bölümü", "Bildiri", "Proje", "Diğer"]:
                                if ptype in grouped:
                                    st.markdown(f"#### {ptype} ({len(grouped[ptype])})")
                                    report_text += f"\n### {ptype}\n\n"
                                    
                                    for idx, pub in enumerate(grouped[ptype], 1):
                                        citation = apa_formatter.format_apa_6(pub)
                                        st.markdown(f"**{idx}.** {citation}")
                                        
                                        # Keep italics for export
                                        report_text += f"{idx}. {citation}\n\n"
                        
                        else:
                            # Simple list for filtered reports
                            if report_type == "Yayın Türü Bazında":
                                st.subheader(f"{selected_pub_type} - {len(filtered_pubs)} Yayın")
                            elif report_type == "Kişi Bazında":
                                st.subheader(f"{selected_person} - {len(filtered_pubs)} Yayın")
                            
                            report_text = ""
                            for idx, pub in enumerate(filtered_pubs, 1):
                                citation = apa_formatter.format_apa_6(pub)
                                ptype = pub.get('publication_type', 'Makale')
                                
                                st.markdown(f"**{idx}. [{ptype}]** {citation}")
                                
                                # Keep italics for export
                                report_text += f"{idx}. [{ptype}] {citation}\n\n"
                        
                        st.markdown("---")
                        st.subheader("📥 Dışa Aktarma")
                        
                        col_exp1, col_exp2 = st.columns(2)
                        
                        with col_exp1:
                            # Word Export
                            try:
                                from docx import Document
                                from docx.shared import Pt, RGBColor, Inches
                                from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
                                import io
                                import re
                                
                                # Create Word document
                                doc = Document()
                                
                                # Add title
                                title_para = doc.add_heading('Akademik Yayın Raporu', 0)
                                title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                
                                # Helper function to add formatted paragraph with italics
                                def add_formatted_paragraph(doc, text, style='Normal'):
                                    """Add paragraph with markdown italics converted to actual italics"""
                                    para = doc.add_paragraph(style=style)
                                    
                                    # Split by italic markers
                                    parts = re.split(r'(\*[^*]+\*)', text)
                                    
                                    for part in parts:
                                        if part.startswith('*') and part.endswith('*'):
                                            # This is italic text
                                            run = para.add_run(part[1:-1])
                                            run.italic = True
                                        elif part:
                                            # Normal text
                                            para.add_run(part)
                                    
                                    return para
                                
                                # Add report content
                                for line in report_text.split('\n'):
                                    if line.strip():
                                        if line.startswith('##'):
                                            doc.add_heading(line.replace('##', '').strip(), level=1)
                                        elif line.startswith('###'):
                                            doc.add_heading(line.replace('###', '').strip(), level=2)
                                        else:
                                            add_formatted_paragraph(doc, line.strip())
                                
                                # Save to bytes
                                docx_buffer = io.BytesIO()
                                doc.save(docx_buffer)
                                docx_buffer.seek(0)
                                
                                st.download_button(
                                    label="📄 Word İndir (.docx)",
                                    data=docx_buffer,
                                    file_name=f"yayin_raporu_{start_date.strftime('%Y%m%d')}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    use_container_width=True
                                )
                            except ImportError:
                                st.warning("Word export için 'python-docx' paketi gerekli. Lütfen yükleyin: pip install python-docx")
                        
                        with col_exp2:
                            # PDF Export with Turkish support
                            try:
                                from reportlab.lib.pagesizes import A4
                                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                                from reportlab.lib.units import cm
                                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                                from reportlab.lib.enums import TA_CENTER
                                from reportlab.pdfbase import pdfmetrics
                                from reportlab.pdfbase.ttfonts import TTFont
                                import io
                                import re
                                import os
                                
                                # Try to register a Turkish-compatible font
                                font_registered = False
                                font_name = 'Helvetica'
                                
                                # Try Windows fonts
                                windows_fonts = [
                                    ('Arial', 'arial.ttf'),
                                    ('Times', 'times.ttf'),
                                    ('Calibri', 'calibri.ttf')
                                ]
                                
                                for fname, ffile in windows_fonts:
                                    try:
                                        font_path = os.path.join('C:\\Windows\\Fonts', ffile)
                                        if os.path.exists(font_path):
                                            pdfmetrics.registerFont(TTFont(fname, font_path))
                                            font_name = fname
                                            font_registered = True
                                            break
                                    except:
                                        continue
                                
                                # Create PDF
                                pdf_buffer = io.BytesIO()
                                doc_pdf = SimpleDocTemplate(pdf_buffer, pagesize=A4, 
                                                           leftMargin=2*cm, rightMargin=2*cm,
                                                           topMargin=2*cm, bottomMargin=2*cm)
                                story = []
                                styles = getSampleStyleSheet()
                                
                                # Custom styles
                                title_style = ParagraphStyle(
                                    'CustomTitle',
                                    parent=styles['Heading1'],
                                    fontSize=18,
                                    fontName=font_name,
                                    textColor='#1f77b4',
                                    spaceAfter=30,
                                    alignment=TA_CENTER
                                )
                                
                                heading1_style = ParagraphStyle(
                                    'CustomH1',
                                    parent=styles['Heading1'],
                                    fontSize=14,
                                    fontName=font_name,
                                    spaceAfter=12
                                )
                                
                                heading2_style = ParagraphStyle(
                                    'CustomH2',
                                    parent=styles['Heading2'],
                                    fontSize=12,
                                    fontName=font_name,
                                    spaceAfter=10
                                )
                                
                                normal_style = ParagraphStyle(
                                    'CustomNormal',
                                    parent=styles['Normal'],
                                    fontSize=10,
                                    fontName=font_name,
                                    spaceAfter=6,
                                    leading=14
                                )
                                
                                # Add title
                                story.append(Paragraph("Akademik Yayın Raporu", title_style))
                                story.append(Spacer(1, 0.5*cm))
                                
                                # Helper to convert markdown italics to HTML and escape special chars
                                def markdown_to_html(text):
                                    """Convert markdown italics to HTML italics and escape XML chars"""
                                    # Escape XML special characters
                                    text = text.replace('&', '&amp;')
                                    text = text.replace('<', '&lt;')
                                    text = text.replace('>', '&gt;')
                                    # Replace *text* with <i>text</i>
                                    text = re.sub(r'\*([^*]+)\*', r'<i>\1</i>', text)
                                    return text
                                
                                # Add content
                                for line in report_text.split('\n'):
                                    if line.strip():
                                        try:
                                            if line.startswith('##'):
                                                clean_line = line.replace('##', '').strip()
                                                story.append(Paragraph(clean_line, heading1_style))
                                            elif line.startswith('###'):
                                                clean_line = line.replace('###', '').strip()
                                                story.append(Paragraph(clean_line, heading2_style))
                                            else:
                                                # Convert markdown italics to HTML
                                                formatted_line = markdown_to_html(line.strip())
                                                story.append(Paragraph(formatted_line, normal_style))
                                            story.append(Spacer(1, 0.2*cm))
                                        except Exception as line_error:
                                            # Skip problematic lines
                                            pass
                                
                                doc_pdf.build(story)
                                pdf_buffer.seek(0)
                                
                                st.download_button(
                                    label="📕 PDF İndir (.pdf)",
                                    data=pdf_buffer,
                                    file_name=f"yayin_raporu_{start_date.strftime('%Y%m%d')}.pdf",
                                    mime="application/pdf",
                                    use_container_width=True
                                )
                            except Exception as e:
                                st.error(f"PDF oluşturulurken hata: {str(e)}")
                                st.warning("PDF export için 'reportlab' paketi gerekli.")

